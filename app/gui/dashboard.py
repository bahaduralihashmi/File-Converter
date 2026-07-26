# app/gui/dashboard.py
"""
Professional Dashboard - Clean Layout with Proper Text Positioning
"""

import os
import sys
import threading
import tempfile
from pathlib import Path
from turtle import color
from typing import List, Dict, Optional
from datetime import datetime

from PySide6.QtCore import Qt, Signal, Slot, QSize, QTimer, QThread
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QGridLayout,
    QLabel, QPushButton, QFrame, QProgressBar,
    QFileDialog, QMessageBox, QScrollArea, QGroupBox,
    QComboBox, QListWidget, QListWidgetItem,
    QLineEdit, QMenu, QApplication, QDialog, QCheckBox,
    QSplitter
)
from PySide6.QtGui import QAction, QFont

from config import Config
from app.core.converter import ConverterEngine
from app.core.file_manager import FileManager
from app.core.history import HistoryManager
from app.core.settings import SettingsManager
from app.core.ai_engine import ai_engine
from app.utils.logger import get_logger
from app.utils.helpers import format_file_size, resource_path
# from dist.AllFilesConverterAI._internal.PySide6.QtGui import QIcon
from PySide6.QtGui import QAction, QFont, QIcon

logger = get_logger(__name__)


# ============================================================
# CONVERSION THREAD
# ============================================================

class ConversionThread(QThread):
    """Background conversion thread"""
    progress = Signal(int, str)
    finished = Signal(bool, int, int, str)  # success, successful, total, message
    file_status = Signal(str, bool)
    error = Signal(str)
    
    def __init__(self, files, format, folder, options, parent=None):
        super().__init__(parent)
        self.files = files
        self.format = format
        self.folder = folder
        self.options = options
        self.is_running = True
        self.converter = ConverterEngine()
        self.history = HistoryManager()
        
    def run(self):
        total = len(self.files)
        successful = 0
        failed = 0
        error_messages = []
        
        for i, file_path in enumerate(self.files):
            if not self.is_running:
                break
                
            progress = int(((i + 1) / total) * 100)
            self.progress.emit(progress, f"Processing {Path(file_path).name}...")
            
            original = Path(file_path)
            save_path = str(Path(self.folder) / f"{original.stem}_converted.{self.format}")
            
            counter = 1
            while Path(save_path).exists():
                save_path = str(Path(self.folder) / f"{original.stem}_converted_{counter}.{self.format}")
                counter += 1
            
            self.progress.emit(progress, f"Converting {Path(file_path).name} to {self.format.upper()}...")
            
            try:
                result = self.converter.convert_file(file_path, self.format, save_path, self.options)
                if result:
                    successful += 1
                    self.history.add_entry(
                        input_file=file_path,
                        output_file=save_path,
                        format=self.format,
                        success=True
                    )
                    self.file_status.emit(file_path, True)
                else:
                    failed += 1
                    error_msg = f"Failed to convert: {Path(file_path).name}"
                    error_messages.append(error_msg)
                    self.history.add_entry(
                        input_file=file_path,
                        output_file=save_path if save_path else "",
                        format=self.format,
                        success=False
                    )
                    self.file_status.emit(file_path, False)
            except Exception as e:
                failed += 1
                error_msg = f"Error converting {Path(file_path).name}: {str(e)}"
                error_messages.append(error_msg)
                logger.error(error_msg)
                self.history.add_entry(
                    input_file=file_path,
                    output_file=save_path if save_path else "",
                    format=self.format,
                    success=False,
                    error=str(e)
                )
                self.file_status.emit(file_path, False)
                self.error.emit(error_msg)
        
        self.progress.emit(100, "Complete!")
        
        # Emit finished with results
        self.finished.emit(
            successful > 0,  # success
            successful,      # successful
            total,           # total
            f"Converted {successful}/{total} files"  # message
        )
    
    def stop(self):
        self.is_running = False
# ============================================================
# PROFESSIONAL DASHBOARD
# ============================================================

class ProfessionalDashboard(QWidget):
    """Professional Dashboard - Clean Layout"""
    
    # Signals
    progress_signal = Signal(int, str)
    conversion_done_signal = Signal(bool, str)
    file_added_signal = Signal(str)
    file_removed_signal = Signal(str)
    status_update_signal = Signal(str)
    voice_command_signal = Signal(str)
    
    _ask_save_location_signal = Signal(str, str, object)
    _update_file_status_signal = Signal(str, bool)
    _show_completion_signal = Signal(int, int, bool, str)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # Core components
        self.converter = ConverterEngine()
        self.file_manager = FileManager()
        self.history = HistoryManager()
        self.settings = SettingsManager.load()
        
        # State
        self.current_files = []
        self.is_converting = False
        self.is_voice_active = False
        self.is_camera_active = False
        self.current_theme = self.settings.get("theme", "professional_blue")
        self.camera_btn = None
        self.conversion_thread = None
        
        # Voice
        self.recognizer = None
        self.microphone = None
        self.tts_engine = None
        
        # Format buttons
        self._format_buttons = []
        self.selected_format = "pdf"
        
        # OCR
        self.ocr_checkbox = None
        self.ocr_lang_combo = None
        
        # Setup signals
        self.progress_signal.connect(self._update_progress_ui)
        self.conversion_done_signal.connect(self._on_conversion_done)
        self.file_added_signal.connect(self._add_file_to_list)
        self.file_removed_signal.connect(self._remove_file_from_list)
        self.status_update_signal.connect(self._update_status)
        self.voice_command_signal.connect(self._process_voice_command)
        
        self._ask_save_location_signal.connect(self._on_ask_save_location)
        self._update_file_status_signal.connect(self._on_update_file_status)
        self._show_completion_signal.connect(self._on_show_completion)

        # Initialize
        self._init_voice()
        self.setup_ui()
        self.setup_connections()
        self.apply_theme(self.current_theme)
        
        QTimer.singleShot(100, self._init_after_startup)
        
        logger.info("Professional Dashboard initialized")
    
    # ============================================================
    # INITIALIZATION
    # ============================================================
    
    def _init_after_startup(self):
        logger.info("Dashboard ready")
    
    def _init_voice(self):
        try:
            import speech_recognition as sr
            self.recognizer = sr.Recognizer()
            self.microphone = sr.Microphone()
            logger.info("Voice recognition initialized")
        except Exception as e:
            logger.warning(f"Voice recognition not available: {e}")
            
        try:
            import pyttsx3
            self.tts_engine = pyttsx3.init()
            self.tts_engine.setProperty('rate', 150)
            self.tts_engine.setProperty('volume', 0.9)
            logger.info("TTS initialized")
        except Exception as e:
            logger.warning(f"TTS not available: {e}")
    
    # ============================================================
    # UI SETUP
    # ============================================================
    
    def setup_ui(self):
        """Setup the complete UI with proper layout"""
        
        # Main layout
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # ===== HEADER =====
        header = self._create_header()
        main_layout.addWidget(header)
        
        # ===== CONTENT =====
        content = self._create_content()
        main_layout.addWidget(content, 1)
        
        # ===== STATUS BAR =====
        status_bar = self._create_status_bar()
        main_layout.addWidget(status_bar)
    
    # ============================================================
    # HEADER
    # ============================================================
    
    def _create_header(self):
        """Create header with icon buttons using QToolBar style"""
        header = QFrame()
        header.setProperty("class", "header")
        header.setFixedHeight(45)
        header.setStyleSheet("""
            QFrame[class="header"] {
                background: #1E3A5F;
                border-bottom: 2px solid #2E5A8A;
            }
        """)
        
        layout = QHBoxLayout(header)
        layout.setContentsMargins(16, 0, 16, 0)
        layout.setSpacing(10)
        
        # Title
        # title = QLabel("📁 All Files Converter AI")
        # title.setStyleSheet("""
        #     font-size: 16px;
        #     font-weight: 600;
        #     color: #FFFFFF;
        #     font-family: 'Segoe UI', sans-serif;
        # """)
        # layout.addWidget(title)
        
        layout.addStretch()
        
        # Search
        self.header_search = QLineEdit()
        self.header_search.setPlaceholderText("Search files...")
        self.header_search.setFixedWidth(200)
        self.header_search.setStyleSheet("""
            QLineEdit {
                background: rgba(255,255,255,0.12);
                border: 1px solid rgba(255,255,255,0.15);
                border-radius: 6px;
                padding: 6px 12px;
                font-size: 12px;
                color: #FFFFFF;
            }
            QLineEdit:focus {
                background: rgba(255,255,255,0.20);
                border-color: rgba(255,255,255,0.30);
            }
            QLineEdit::placeholder {
                color: rgba(255,255,255,0.50);
            }
        """)
        self.header_search.textChanged.connect(self.filter_files)
        layout.addWidget(self.header_search)
        
        # ===== ICON BUTTONS WITH SVG FALLBACK =====
        # Use emoji with colored backgrounds - they WILL render
        actions = [
            ("🎤", "Voice Control", self.toggle_voice),
            ("📷", "Camera Scanner", self.toggle_camera),
            ("⚙️", "Settings", self.open_settings),
        ]
        
        for emoji, tip, callback in actions:
            btn = QPushButton(emoji)
            btn.setFixedSize(50, 50)
            btn.setToolTip(tip)
            btn.setCursor(Qt.PointingHandCursor)
            btn.setStyleSheet("""
                QPushButton {
                    border: none;
                    border-radius: 3px;
                    background: rgba(255,255,255,0.10);
                    color: #FFFFFF;
                    font-size: 18px;
                    font-weight: normal;
                }
                QPushButton:hover {
                    background: rgba(255,255,255,0.20);
                }
                QPushButton:pressed {
                    background: rgba(255,255,255,0.30);
                }
            """)
            btn.clicked.connect(callback)
            
            if emoji == "📷":
                self.camera_btn = btn
            
            layout.addWidget(btn)
        
        return header
    
    def open_settings(self):
        main_window = self.window()
        if hasattr(main_window, 'switch_page'):
            main_window.switch_page('settings')
    
    # ============================================================
    # CONTENT
    # ============================================================
    
    def _create_content(self):
        """Create content with proper spacing"""
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(12)
        
        # Stats Row
        layout.addLayout(self._create_stats_row())
        
        # Two column layout
        content_row = QHBoxLayout()
        content_row.setSpacing(16)
        
        # Left Column (File Management)
        left_widget = self._create_left_column()
        content_row.addWidget(left_widget, 3)
        
        # Right Column (Controls)
        right_widget = self._create_right_column()
        content_row.addWidget(right_widget, 2)
        
        layout.addLayout(content_row)
        layout.addStretch()
        
        return container
    
    # ============================================================
    # STATS ROW
    # ============================================================
    
    def _create_stats_row(self):
        """Create statistics row"""
        layout = QHBoxLayout()
        layout.setSpacing(10)
        
        stats = [
            ("📄", "Files", "0"),
            ("🔄", "Today", "0"),
            ("📋", "Queue", "0"),
            ("✅", "Rate", "100%"),
        ]
        
        self.stats_cards = []
        for icon, label, value in stats:
            card = QFrame()
            card.setFixedHeight(44)
            card.setStyleSheet("""
                QFrame {
                    background: #FFFFFF;
                    border: 1px solid #E8E8E8;
                    border-radius: 8px;
                }
            """)
            
            card_layout = QHBoxLayout(card)
            card_layout.setContentsMargins(12, 4, 12, 4)
            card_layout.setSpacing(8)
            
            icon_label = QLabel(icon)
            icon_label.setStyleSheet("font-size: 18px;")
            icon_label.setFixedWidth(28)
            card_layout.addWidget(icon_label)
            
            text_widget = QWidget()
            text_layout = QVBoxLayout(text_widget)
            text_layout.setContentsMargins(0, 0, 0, 0)
            text_layout.setSpacing(0)
            
            value_label = QLabel(value)
            value_label.setStyleSheet("font-size: 14px; font-weight: 700; color: #1A202C;")
            text_layout.addWidget(value_label)
            
            label_label = QLabel(label)
            label_label.setStyleSheet("font-size: 9px; color: #718096; font-weight: 500;")
            text_layout.addWidget(label_label)
            
            card_layout.addWidget(text_widget)
            card_layout.addStretch()
            
            layout.addWidget(card)
            card.value_label = value_label
            self.stats_cards.append(card)
        
        return layout
    
    # ============================================================
    # LEFT COLUMN - FILE MANAGEMENT
    # ============================================================
    
    def _create_left_column(self):
        """Create left column with file management"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)
        
        # Drop Zone
        drop = self._create_drop_zone()
        layout.addWidget(drop)
        
        # File List
        file_list = self._create_file_list()
        layout.addWidget(file_list)
        
        return widget
    
    def _create_drop_zone(self):
        """Create drop zone"""
        drop = QFrame()
        drop.setMinimumHeight(60)
        drop.setAcceptDrops(True)
        drop.setStyleSheet("""
            QFrame {
                background: #F7FAFC;
                border: 2px dashed #CBD5E0;
                border-radius: 8px;
            }
            QFrame:hover {
                background: #EDF2F7;
                border-color: #2E5A8A;
            }
        """)
        
        drop_layout = QVBoxLayout(drop)
        drop_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        drop_layout.setSpacing(2)
        
        icon = QLabel("📥")
        icon.setStyleSheet("font-size: 24px;")
        icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        drop_layout.addWidget(icon)
        
        title = QLabel("Drop files here or click to browse")
        title.setStyleSheet("font-size: 12px; font-weight: 500; color: #2D3748;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        drop_layout.addWidget(title)
        
        subtitle = QLabel("Supports 1000+ formats")
        subtitle.setStyleSheet("font-size: 10px; color: #A0AEC0;")
        subtitle.setAlignment(Qt.AlignmentFlag.AlignCenter)
        drop_layout.addWidget(subtitle)
        
        drop.mousePressEvent = lambda e: self.add_files()
        drop.dragEnterEvent = self._drag_enter
        drop.dragLeaveEvent = self._drag_leave
        drop.dropEvent = self._drop_event
        
        return drop
    
    def _create_file_list(self):
        """Create file list with controls"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(4)
        
        # Header
        header = QHBoxLayout()
        header.setSpacing(8)
        
        title = QLabel("📋 File List")
        title.setStyleSheet("font-size: 12px; font-weight: 600; color: #1A202C;")
        header.addWidget(title)
        header.addStretch()
        
        self.file_count_label = QLabel("0 files")
        self.file_count_label.setStyleSheet("font-size: 10px; color: #718096;")
        header.addWidget(self.file_count_label)
        layout.addLayout(header)
        
        # List
        self.file_list = QListWidget()
        self.file_list.setMinimumHeight(100)
        self.file_list.setMaximumHeight(130)
        self.file_list.setSelectionMode(QListWidget.ExtendedSelection)
        self.file_list.setStyleSheet("""
            QListWidget {
                border: 1px solid #E2E8F0;
                border-radius: 6px;
                padding: 4px;
                background: #FFFFFF;
            }
            QListWidget::item {
                padding: 4px 8px;
                border-radius: 4px;
                color: #1A202C;
                font-size: 11px;
            }
            QListWidget::item:hover {
                background: #EDF2F7;
            }
            QListWidget::item:selected {
                background: #2E5A8A;
                color: white;
            }
        """)
        self.file_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.file_list.customContextMenuRequested.connect(self._show_context_menu)
        layout.addWidget(self.file_list)
        
        # Buttons
        btn_row = QHBoxLayout()
        btn_row.setSpacing(4)
        
        btn_style = """
            QPushButton {
                padding: 4px 12px;
                border: none;
                border-radius: 4px;
                font-size: 10px;
                font-weight: 500;
                min-height: 24px;
            }
        """
        
        add_btn = QPushButton("📂 Add Files")
        add_btn.setStyleSheet(btn_style + "background: #2E5A8A; color: black;")
        add_btn.clicked.connect(self.add_files)
        btn_row.addWidget(add_btn)
        
        add_folder_btn = QPushButton("📁 Add Folder")
        add_folder_btn.setStyleSheet(btn_style + "background: #EDF2F7; color: #2D3748;")
        add_folder_btn.clicked.connect(self.add_folder)
        btn_row.addWidget(add_folder_btn)
        
        remove_btn = QPushButton("🗑 Remove")
        remove_btn.setStyleSheet(btn_style + "background: #FED7D7; color: #C53030;")
        remove_btn.clicked.connect(self.remove_selected_files)
        btn_row.addWidget(remove_btn)
        
        clear_btn = QPushButton("Clear All")
        clear_btn.setStyleSheet(btn_style + "background: #FED7D7; color: #C53030;")
        clear_btn.clicked.connect(self.remove_all_files)
        btn_row.addWidget(clear_btn)
        
        btn_row.addStretch()
        layout.addLayout(btn_row)
        
        return widget
    
    # ============================================================
    # RIGHT COLUMN - CONTROLS
    # ============================================================
    
    def _create_right_column(self):
        """Create right column with controls"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)
        
        # Conversion Controls
        controls = self._create_conversion_controls()
        layout.addWidget(controls)
        
        # Quick Actions
        quick_actions = self._create_quick_actions()
        layout.addWidget(quick_actions)
        
        # AI Assistant
        ai_button = self._create_ai_button()
        layout.addWidget(ai_button)
        
        # Supported Formats
        categories = self._create_categories()
        layout.addWidget(categories)
        
        layout.addStretch()
        return widget
    
    # ============================================================
    # CONVERSION CONTROLS - ALL 104 FORMATS
    # ============================================================
    
    def _create_conversion_controls(self):
        """Create conversion controls with ALL formats"""
        group = QGroupBox("Conversion Controls")
        group.setStyleSheet("""
            QGroupBox {
                border: 1px solid #E2E8F0;
                border-radius: 8px;
                padding-top: 10px;
                font-size: 11px;
                font-weight: 600;
                color: #1A202C;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 6px;
            }
        """)
        
        layout = QVBoxLayout(group)
        layout.setContentsMargins(12, 4, 12, 8)
        layout.setSpacing(4)
        
        # Format selection
        format_row = QHBoxLayout()
        format_row.setSpacing(6)
        
        format_label = QLabel("Output:")
        format_label.setStyleSheet("font-size: 11px; font-weight: 500; color: #2D3748;")
        format_row.addWidget(format_label)
        format_row.addStretch()
        
        self.format_combo = QComboBox()
        self.format_combo.setMinimumWidth(160)
        self.format_combo.setStyleSheet("""
            QComboBox {
                padding: 4px 8px;
                border: 1px solid #E2E8F0;
                border-radius: 4px;
                font-size: 11px;
                min-width: 140px;
                background: #FFFFFF;
            }
            QComboBox::drop-down {
                border: none;
                padding-right: 4px;
            }
        """)
        
        # ===== ALL 104 FORMATS =====
        # Documents (20)
        for fmt in ['pdf', 'docx', 'doc', 'txt', 'rtf', 'odt', 'html', 'md', 
                    'tex', 'pages', 'key', 'numbers', 'xml', 'json', 'csv', 
                    'xlsx', 'xls', 'pptx', 'ppt', 'mht']:
            self.format_combo.addItem(f"📄 {fmt.upper()}", fmt)
        
        # Images (21)
        for fmt in ['jpg', 'jpeg', 'png', 'webp', 'bmp', 'gif', 'tiff', 'tif',
                    'ico', 'svg', 'heic', 'avif', 'raw', 'cr2', 'nef', 'arw',
                    'dng', 'orf', 'rw2', 'pef', 'srw']:
            self.format_combo.addItem(f"🖼 {fmt.upper()}", fmt)
        
        # Audio (9)
        for fmt in ['mp3', 'wav', 'flac', 'aac', 'ogg', 'wma', 'm4a', 'aiff', 'opus']:
            self.format_combo.addItem(f"🎵 {fmt.upper()}", fmt)
        
        # Video (12)
        for fmt in ['mp4', 'avi', 'mkv', 'mov', 'wmv', 'flv', 'webm', 
                    'm4v', '3gp', 'mpg', 'mpeg', 'm2ts']:
            self.format_combo.addItem(f"🎬 {fmt.upper()}", fmt)
        
        # Archives (14)
        for fmt in ['zip', 'rar', '7z', 'tar', 'gz', 'bz2', 'xz', 
                    'tgz', 'tbz2', 'txz', 'zst', 'iso', 'cab', 'lzma']:
            self.format_combo.addItem(f"📦 {fmt.upper()}", fmt)
        
        # E-books (18)
        for fmt in ['epub', 'mobi', 'azw3', 'fb2', 'lit', 'lrf', 'azw', 
                    'kfx', 'pdb', 'prc', 'pml', 'rb', 'snb', 'cbr', 'cbz', 
                    'pdf', 'docx', 'txt']:
            self.format_combo.addItem(f"📚 {fmt.upper()}", fmt)
        
        # Spreadsheets (10)
        for fmt in ['xlsx', 'xls', 'csv', 'ods', 'tsv', 'numbers', 
                    'html', 'json', 'xml', 'dif']:
            self.format_combo.addItem(f"📊 {fmt.upper()}", fmt)
        
        self.format_combo.setCurrentIndex(0)
        format_row.addWidget(self.format_combo)
        layout.addLayout(format_row)
        
        # Quick formats
        quick_row = QHBoxLayout()
        quick_row.setSpacing(4)
        
        for label, fmt in [("PDF", "pdf"), ("DOCX", "docx"), ("JPG", "jpg"), 
                          ("PNG", "png"), ("MP3", "mp3"), ("MP4", "mp4"), 
                          ("ZIP", "zip"), ("EPUB", "epub")]:
            btn = QPushButton(label)
            btn.setStyleSheet("""
                QPushButton {
                    padding: 2px 4px;
                    border: 1px solid #E2E8F0;
                    border-radius: 3px;
                    font-size: 8px;
                    background: #F7FAFC;
                    color: #2D3748;
                    min-width: 28px;
                }
                QPushButton:hover {
                    background: #EDF2F7;
                }
            """)
            btn.clicked.connect(lambda checked, f=fmt: self.quick_convert(f))
            quick_row.addWidget(btn)
        quick_row.addStretch()
        layout.addLayout(quick_row)
        
        # Convert button
        self.convert_btn = QPushButton("🔄 Convert Now")
        self.convert_btn.setStyleSheet("""
            QPushButton {
                padding: 6px 16px;
                border: none;
                border-radius: 4px;
                background: #2E5A8A;
                color: white;
                font-size: 12px;
                font-weight: 600;
                min-height: 28px;
            }
            QPushButton:hover {
                background: #1E3A5F;
            }
            QPushButton:disabled {
                background: #A0AEC0;
            }
        """)
        self.convert_btn.clicked.connect(self.start_conversion)
        layout.addWidget(self.convert_btn)
        
        # Options
        options_row = QHBoxLayout()
        options_row.setSpacing(8)
        
        self.ask_save_location = QCheckBox("Ask where to save")
        self.ask_save_location.setChecked(True)
        self.ask_save_location.setStyleSheet("font-size: 10px; color: #2D3748;")
        options_row.addWidget(self.ask_save_location)
        options_row.addStretch()
        layout.addLayout(options_row)
        
        # Progress
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: none;
                border-radius: 3px;
                background: #E2E8F0;
                height: 4px;
            }
            QProgressBar::chunk {
                background: #2E5A8A;
                border-radius: 3px;
            }
        """)
        layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel("Ready")
        self.status_label.setStyleSheet("font-size: 10px; color: #718096; padding: 2px 0;")
        layout.addWidget(self.status_label)
        
        # Format count
        count_label = QLabel("📊 104+ Formats Available")
        count_label.setStyleSheet("font-size: 9px; color: #A0AEC0; padding-top: 2px;")
        count_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(count_label)
        
        return group
    
    # ============================================================
    # QUICK ACTIONS
    # ============================================================
    
    def _create_quick_actions(self):
        """Create quick actions"""
        group = QGroupBox("Quick Actions")
        group.setStyleSheet("""
            QGroupBox {
                border: 1px solid #E2E8F0;
                border-radius: 8px;
                padding-top: 10px;
                font-size: 11px;
                font-weight: 600;
                color: #1A202C;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 6px;
            }
        """)
        
        layout = QVBoxLayout(group)
        layout.setContentsMargins(12, 4, 12, 8)
        layout.setSpacing(4)
        
        btn = QPushButton("📄 OCR Extract Text")
        btn.setStyleSheet("""
            QPushButton {
                padding: 4px 12px;
                border: 1px solid #E2E8F0;
                border-radius: 4px;
                background: #F7FAFC;
                color: #2D3748;
                font-size: 11px;
                text-align: left;
            }
            QPushButton:hover {
                background: #EDF2F7;
            }
        """)
        btn.clicked.connect(self.open_ocr)
        layout.addWidget(btn)
        
        return group
    
    # ============================================================
    # AI ASSISTANT
    # ============================================================
    
    def _create_ai_button(self):
        """Create AI Assistant button"""
        group = QGroupBox("AI Assistant")
        group.setStyleSheet("""
            QGroupBox {
                border: 1px solid #E2E8F0;
                border-radius: 8px;
                padding-top: 10px;
                font-size: 11px;
                font-weight: 600;
                color: #1A202C;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 6px;
            }
        """)
        
        layout = QVBoxLayout(group)
        layout.setContentsMargins(12, 4, 12, 8)
        layout.setSpacing(4)
        
        self.ai_popup_btn = QPushButton("🤖 Open AI Assistant")
        self.ai_popup_btn.setStyleSheet("""
            QPushButton {
                padding: 6px 16px;
                border: none;
                border-radius: 4px;
                background: #2E5A8A;
                color: white;
                font-size: 12px;
                font-weight: 600;
                min-height: 28px;
            }
            QPushButton:hover {
                background: #1E3A5F;
            }
        """)
        self.ai_popup_btn.clicked.connect(self.open_ai_popup)
        layout.addWidget(self.ai_popup_btn)
        
        self.ai_status_label = QLabel("✅ AI Ready")
        self.ai_status_label.setStyleSheet("font-size: 10px; color: #718096;")
        self.ai_status_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.ai_status_label)
        
        return group
    
    # ============================================================
    # SUPPORTED FORMATS
    # ============================================================
    
    def _create_categories(self):
        """Create categories display"""
        group = QGroupBox("Supported Formats")
        group.setStyleSheet("""
            QGroupBox {
                border: 1px solid #E2E8F0;
                border-radius: 8px;
                padding-top: 10px;
                font-size: 11px;
                font-weight: 600;
                color: #1A202C;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 6px;
            }
        """)
        
        layout = QVBoxLayout(group)
        layout.setContentsMargins(12, 4, 12, 8)
        layout.setSpacing(2)
        
        # categories = [
        #     ("📄 Documents", "20"),
        #     ("🖼 Images", "21"),
        #     ("🎵 Audio", "9"),
        #     ("🎬 Video", "12"),
        #     ("📦 Archives", "14"),
        #     ("📚 E-books", "18"),
        #     ("📊 Spreadsheets", "10"),
        # ]
        
        # for cat, count in categories:
        #     row = QHBoxLayout()
        #     row.setSpacing(4)
            
        #     cat_label = QLabel(cat)
        #     cat_label.setStyleSheet("font-size: 10px; font-weight: 500; color: #2D3748;")
        #     row.addWidget(cat_label)
        #     row.addStretch()
            
        #     count_label = QLabel(count)
        #     count_label.setStyleSheet("""
        #         font-size: 8px;
        #         color: #718096;
        #         background: #EDF2F7;
        #         padding: 0 8px;
        #         border-radius: 8px;
        #         font-weight: 600;
        #     """)
        #     row.addWidget(count_label)
        #     layout.addLayout(row)
        
        total_row = QHBoxLayout()
        total_row.addStretch()
        total_label = QLabel("✨ 104+ Formats Total")
        total_label.setStyleSheet("""
            font-size: 9px;
            color: #2E5A8A;
            font-weight: 700;
            padding: 2px 0;
        """)
        total_row.addWidget(total_label)
        layout.addLayout(total_row)
        
        return group
    
    # ============================================================
    # STATUS BAR
    # ============================================================
    
    def _create_status_bar(self):
        """Create status bar"""
        status = QFrame()
        status.setFixedHeight(24)
        status.setStyleSheet("""
            QFrame {
                background: #F7FAFC;
                border-top: 1px solid #E2E8F0;
                padding: 0 16px;
            }
        """)
        
        layout = QHBoxLayout(status)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)
        
        self.status_label_bar = QLabel("🟢 Ready")
        self.status_label_bar.setStyleSheet("font-size: 10px; color: #718096;")
        layout.addWidget(self.status_label_bar)
        
        layout.addStretch()
        
        time_label = QLabel(datetime.now().strftime("%I:%M %p"))
        time_label.setStyleSheet("font-size: 10px; color: #718096;")
        layout.addWidget(time_label)
        
        version = QLabel(f"v{Config.APP_VERSION}")
        version.setStyleSheet("font-size: 10px; color: #718096;")
        layout.addWidget(version)
        
        return status
    
    # ============================================================
    # DRAG AND DROP
    # ============================================================
    
    def _drag_enter(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
    
    def _drag_leave(self, event):
        pass
    
    def _drop_event(self, event):
        files = []
        for url in event.mimeData().urls():
            file_path = url.toLocalFile()
            if file_path and Path(file_path).exists():
                files.append(file_path)
        if files:
            for file_path in files:
                self.file_added_signal.emit(file_path)
    
    # ============================================================
    # CONTEXT MENU
    # ============================================================
    
    def _show_context_menu(self, position):
        menu = QMenu()
        menu.setStyleSheet("""
            QMenu {
                background: #FFFFFF;
                border: 1px solid #E2E8F0;
                border-radius: 6px;
                padding: 4px;
            }
            QMenu::item {
                padding: 4px 16px;
                font-size: 11px;
                color: #2D3748;
            }
            QMenu::item:selected {
                background: #EDF2F7;
            }
        """)
        
        remove_action = QAction("🗑 Remove Selected", self)
        remove_action.triggered.connect(self.remove_selected_file)
        menu.addAction(remove_action)
        
        remove_all_action = QAction("🗑 Remove All", self)
        remove_all_action.triggered.connect(self.remove_all_files)
        menu.addAction(remove_all_action)
        
        menu.exec_(self.file_list.mapToGlobal(position))
    
    def remove_selected_file(self):
        current_row = self.file_list.currentRow()
        if current_row >= 0:
            item = self.file_list.item(current_row)
            if item:
                file_path = item.data(Qt.UserRole)
                self.file_removed_signal.emit(file_path)
    
    # ============================================================
    # FILE OPERATIONS
    # ============================================================
    
    def add_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, 
            "Select Files to Convert", 
            "",
            "All Files (*.*)"
        )
        if files:
            added_count = 0
            for file_path in files:
                if file_path not in self.current_files:
                    self._add_file_to_list(file_path)
                    added_count += 1
            self._update_status(f"✅ Added {added_count} file(s)")
    
    def add_folder(self):
        folder = QFileDialog.getExistingDirectory(
            self,
            "Select Folder to Add Files",
            "",
            QFileDialog.ShowDirsOnly
        )
        if folder:
            folder_path = Path(folder)
            added_count = 0
            for file_path in folder_path.iterdir():
                if file_path.is_file() and str(file_path) not in self.current_files:
                    self._add_file_to_list(str(file_path))
                    added_count += 1
            self._update_status(f"✅ Added {added_count} files from {folder_path.name}")
    
    def remove_all_files(self):
        if self.current_files:
            reply = QMessageBox.question(
                self,
                "Remove All Files",
                "Remove all files from the list?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self.file_list.clear()
                self.current_files = []
                self.update_file_count()
                self._update_status("All files removed")
    
    def remove_selected_files(self):
        selected_items = self.file_list.selectedItems()
        if not selected_items:
            QMessageBox.information(self, "No Selection", "Please select files to remove.")
            return
        
        count = len(selected_items)
        reply = QMessageBox.question(
            self,
            "Remove Files",
            f"Remove {count} selected file(s)?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            for item in selected_items:
                file_path = item.data(Qt.UserRole)
                if file_path in self.current_files:
                    self.current_files.remove(file_path)
                    self.file_list.takeItem(self.file_list.row(item))
            self.update_file_count()
            self._update_status(f"🗑️ Removed {count} file(s)")
    
    # ============================================================
    # FILE LIST MANAGEMENT
    # ============================================================
    
    @Slot(str)
    def _add_file_to_list(self, file_path):
        if file_path not in self.current_files:
            self.current_files.append(file_path)
            path = Path(file_path)
            if not path.exists():
                logger.warning(f"File does not exist: {file_path}")
                return
            size = format_file_size(path.stat().st_size)
            item = QListWidgetItem(f"📄 {path.name} ({size})")
            item.setData(Qt.UserRole, file_path)
            self.file_list.addItem(item)
            self.update_file_count()
            logger.info(f"File added to list: {file_path}")
    
    @Slot(str)
    def _remove_file_from_list(self, file_path):
        if file_path in self.current_files:
            self.current_files.remove(file_path)
            for i in range(self.file_list.count()):
                item = self.file_list.item(i)
                if item.data(Qt.UserRole) == file_path:
                    self.file_list.takeItem(i)
                    break
            self.update_file_count()
    
    def update_file_count(self):
        total = len(self.current_files)
        self.file_count_label.setText(f"{total} files")
        if hasattr(self, 'stats_cards') and len(self.stats_cards) > 0:
            self.stats_cards[0].value_label.setText(str(total))
            self.stats_cards[2].value_label.setText(str(total))
    
    # ============================================================
    # SEARCH
    # ============================================================
    
    def filter_files(self, search_text: str):
        """Filter files in the list based on search text"""
        search_text = search_text.lower().strip()
        
        if not search_text:
            for i in range(self.file_list.count()):
                item = self.file_list.item(i)
                item.setHidden(False)
            return
        
        visible_count = 0
        for i in range(self.file_list.count()):
            item = self.file_list.item(i)
            file_name = item.text().lower()
            if search_text in file_name:
                item.setHidden(False)
                visible_count += 1
            else:
                item.setHidden(True)
        
        self.status_label.setText(f"🔍 Found {visible_count} files")
        self.status_label_bar.setText(f"🔍 Found {visible_count} of {len(self.current_files)} files")
    
    # ============================================================
    # STATUS UPDATES
    # ============================================================
    
    @Slot(str)
    def _update_status(self, message):
        self.status_label.setText(message)
        self.status_label_bar.setText(message)
    
    @Slot(int, str)
    def _update_progress_ui(self, progress, message):
        self.progress_bar.setValue(progress)
        self.status_label.setText(message)
        self.status_label_bar.setText(message)
    
    # ============================================================
    # CONVERSION
    # ============================================================
    
    def get_conversion_options(self) -> dict:
        options = {}
        quality_map = {"low": 50, "medium": 75, "high": 95}
        quality = self.settings.get("quality", "high")
        options['quality'] = quality_map.get(quality, 95)
        options['keep_original'] = self.settings.get("keep_original", True)
        options['ask_location'] = self.ask_save_location.isChecked()
        return options
    
    def quick_convert(self, format_name):
        index = self.format_combo.findData(format_name)
        if index >= 0:
            self.format_combo.setCurrentIndex(index)
        self.start_conversion()
    
    def start_conversion(self):
        """Start conversion - Ask for location once for multiple files"""
        if not self.current_files:
            QMessageBox.warning(self, "No Files", "Please add files to convert!")
            return
        
        if self.is_converting:
            QMessageBox.warning(self, "Busy", "Conversion already in progress!")
            return
        
        # Get valid files
        valid_files = []
        for f in self.current_files:
            if Path(f).exists():
                valid_files.append(f)
            else:
                logger.warning(f"File does not exist: {f}")
        
        if not valid_files:
            QMessageBox.warning(self, "No Files", "No valid files to convert!")
            return
        
        ask_each = self.ask_save_location.isChecked()
        output_folder_path = None
        
        if ask_each:
            output_folder = self.settings.get("output_folder", str(Path.home() / "Documents"))
            folder = QFileDialog.getExistingDirectory(
                self,
                "Select Output Folder for All Files",
                output_folder,
                QFileDialog.ShowDirsOnly
            )
            if not folder:
                return
            
            self.settings.set("output_folder", folder)
            self.settings.save()
            output_folder_path = folder
        else:
            output_folder = self.settings.get("output_folder", str(Path.home() / "Documents"))
            if not Path(output_folder).exists():
                output_folder = str(Path.home() / "Documents")
                self.settings.set("output_folder", output_folder)
                self.settings.save()
            output_folder_path = output_folder
        
        self.is_converting = True
        self.convert_btn.setEnabled(False)
        self.convert_btn.setText("⏳ Converting...")
        self._update_status(f"Converting {len(valid_files)} files...")
        
        output_format = self.format_combo.currentData()
        options = self.get_conversion_options()
        
        # ===== Create thread =====
        self.conversion_thread = ConversionThread(
            valid_files, 
            output_format, 
            output_folder_path, 
            options,
            self
        )
        
        # ===== Connect signals =====
        self.conversion_thread.progress.connect(self._on_conversion_progress)
        self.conversion_thread.finished.connect(self._on_conversion_finished)
        self.conversion_thread.file_status.connect(self._on_file_status)
        self.conversion_thread.error.connect(self._on_conversion_error)
        
        # Start thread
        self.conversion_thread.start()
    
    def _process_conversion(self, files, output_format, output_folder, options):
        """Process conversion in background - All files to one folder"""
        total = len(files)
        successful = 0
        failed = 0
        
        for i, file_path in enumerate(files):
            if not self.is_converting:
                break
            
            progress = int(((i + 1) / total) * 100)
            self.progress_signal.emit(progress, f"Processing {Path(file_path).name}...")
            
            # ===== FIX: Generate output path in the selected folder =====
            original = Path(file_path)
            save_path = str(Path(output_folder) / f"{original.stem}_converted.{output_format}")
            
            # Ensure unique filename
            counter = 1
            while Path(save_path).exists():
                save_path = str(Path(output_folder) / f"{original.stem}_converted_{counter}.{output_format}")
                counter += 1
            
            self.progress_signal.emit(progress, f"Converting {Path(file_path).name} to {output_format.upper()}...")
            
            try:
                result = self.converter.convert_file(file_path, output_format, save_path, options)
                if result:
                    successful += 1
                    self.status_update_signal.emit(f"✅ Saved: {Path(save_path).name}")
                    self.history.add_entry(
                        input_file=file_path,
                        output_file=save_path,
                        format=output_format,
                        success=True
                    )
                else:
                    failed += 1
                    self.status_update_signal.emit(f"❌ Failed: {Path(file_path).name}")
                    self.history.add_entry(
                        input_file=file_path,
                        output_file=save_path if save_path else "",
                        format=output_format,
                        success=False
                    )
                self._update_file_status_signal.emit(file_path, result)
            except Exception as e:
                failed += 1
                logger.error(f"Conversion error: {e}")
                self._update_file_status_signal.emit(file_path, False)
                self.status_update_signal.emit(f"❌ Error: {str(e)[:50]}")
                self.history.add_entry(
                    input_file=file_path,
                    output_file=save_path if save_path else "",
                    format=output_format,
                    success=False,
                    error=str(e)
                )
        
        self.is_converting = False
        self.convert_btn.setEnabled(True)
        self.convert_btn.setText("🔄 Convert Now")
        self.progress_signal.emit(100, "Complete!")
        
        # ===== Single completion message =====
        if successful > 0:
            if hasattr(self, 'stats_cards') and len(self.stats_cards) > 1:
                current = int(self.stats_cards[1].value_label.text()) if self.stats_cards[1].value_label.text().isdigit() else 0
                self.stats_cards[1].value_label.setText(str(current + successful))
            
            self.status_update_signal.emit(f"✅ Converted {successful} files")
            
            if failed > 0:
                msg = f"✅ Converted {successful}/{total} files successfully!\n❌ Failed: {failed} file(s)\n\n📁 Location: {output_folder}"
                QMessageBox.information(self, "Conversion Complete", msg)
            else:
                QMessageBox.information(self, "Conversion Complete", f"✅ All {successful} files converted successfully!\n\n📁 Location: {output_folder}")
        else:
            self.status_update_signal.emit("No files converted")
            QMessageBox.warning(self, "Conversion Failed", "No files were converted. Please check your files and try again.")
        
        # Refresh history
        QTimer.singleShot(500, self._refresh_history_view)
    
    # ============================================================
    # SIGNAL HANDLERS
    # ============================================================
    
    @Slot(bool, str)
    def _on_conversion_done(self, success, message):
        self.is_converting = False
        self.convert_btn.setEnabled(True)
        self.convert_btn.setText("🔄 Convert Now")
        self.status_label.setText(message)
        self.status_label_bar.setText(message)
        
        if success:
            QMessageBox.information(self, "Conversion Complete", f"✅ {message}")
        else:
            QMessageBox.warning(self, "Conversion Failed", f"❌ {message}")
        
        self._refresh_history_view()
    
    @Slot(str, str, object)
    def _on_ask_save_location(self, file_path, output_format, result_queue):
        original = Path(file_path)
        default_name = f"{original.stem}_converted.{output_format}"
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Converted File As",
            str(original.parent / default_name),
            f"{output_format.upper()} Files (*.{output_format});;All Files (*.*)"
        )
        
        if file_path:
            path = Path(file_path)
            if path.suffix.lower() != f".{output_format}":
                file_path = str(path.parent / f"{path.stem}.{output_format}")
            result_queue.put(file_path)
        else:
            result_queue.put(None)
    
    @Slot(str, bool)
    def _on_update_file_status(self, file_path, success):
        for idx in range(self.file_list.count()):
            item = self.file_list.item(idx)
            if item.data(Qt.UserRole) == file_path:
                status = "✅ Complete" if success else "❌ Failed"
                current_text = item.text()
                if "(" in current_text and ")" in current_text:
                    parts = current_text.split(" (")
                    if len(parts) >= 2:
                        current_text = " (".join(parts[:-1]) + ")"
                item.setText(f"{current_text} ({status})")
                break
    
    @Slot(int, int, bool, str)
    def _on_show_completion(self, successful, total, ask_each, output_folder):
        if ask_each:
            msg = f"✅ Converted {successful}/{total} files successfully!"
        else:
            msg = f"✅ Converted {successful}/{total} files successfully!\n\nOutput folder: {output_folder}"
        QMessageBox.information(self, "Conversion Complete", msg)
    
    def _refresh_history_view(self):
        try:
            main_window = self.window()
            if hasattr(main_window, 'history_page') and main_window.history_page:
                main_window.history_page.load_history()
        except Exception as e:
            logger.error(f"Failed to refresh history: {e}")
    
    # ============================================================
    # THEME
    # ============================================================
    
    def apply_theme(self, theme_name: str):
        """Apply a theme to the entire application"""
        from app.gui.themes import ThemeManager
        
        self.current_theme = theme_name
        self.settings.set("theme", theme_name)
        self.settings.save()
        
        # Apply theme using the class method
        ThemeManager.apply_theme(theme_name)
        
        self.status_update_signal.emit(f"Theme applied: {theme_name}")
        logger.info(f"Theme applied: {theme_name}")
    # ============================================================
    # SETUP CONNECTIONS
    # ============================================================
    
    def setup_connections(self):
        if hasattr(self.converter, 'progress_updated'):
            self.converter.progress_updated.connect(self._on_converter_progress)
        if hasattr(self.converter, 'conversion_complete'):
            self.converter.conversion_complete.connect(self._on_converter_complete)


    @Slot(int, str)
    def _on_conversion_progress(self, progress, message):
        """Handle conversion progress - UI thread"""
        self.progress_bar.setValue(progress)
        self.status_label.setText(message)
        self.status_label_bar.setText(message)

    @Slot(bool, int, int, str)
    def _on_conversion_finished(self, success, successful, total, message):
        """Handle conversion finished - UI thread"""
        self.is_converting = False
        self.convert_btn.setEnabled(True)
        self.convert_btn.setText("🔄 Convert Now")
        self.progress_bar.setValue(100)
        self.status_label.setText(message)
        self.status_label_bar.setText(message)
        
        # Update stats
        if hasattr(self, 'stats_cards') and self.stats_cards and len(self.stats_cards) > 1:
            try:
                current_text = self.stats_cards[1].value_label.text()
                current = int(current_text) if current_text.isdigit() else 0
                self.stats_cards[1].value_label.setText(str(current + successful))
            except (ValueError, AttributeError):
                pass
        
        # Refresh history
        QTimer.singleShot(500, self._refresh_history_view)
        
        # Show dialog in UI thread
        if success:
            if successful == total:
                QMessageBox.information(
                    self, 
                    "Conversion Complete", 
                    f"✅ All {successful} files converted successfully!"
                )
            else:
                QMessageBox.information(
                    self, 
                    "Conversion Complete", 
                    f"✅ Converted {successful}/{total} files successfully!"
                )
        else:
            QMessageBox.warning(
                self, 
                "Conversion Failed", 
                f"❌ No files were converted.\n\n{message}"
            )
        
    @Slot(str)
    def _on_conversion_error(self, error_msg):
        """Handle conversion errors - UI thread"""
        logger.error(error_msg)
        # Optionally show errors in status
        self.status_label.setText(f"⚠️ {error_msg[:100]}")
        
    @Slot(list)
    def _on_converter_complete(self, results):
        self.conversion_done_signal.emit(True, "Conversion complete")

    @Slot(str, bool)
    def _on_file_status(self, file_path, success):
        """Handle individual file status - UI thread"""
        from PySide6.QtGui import QColor
        
        for idx in range(self.file_list.count()):
            item = self.file_list.item(idx)
            if item.data(Qt.UserRole) == file_path:
                if success:
                    status = "✅ Complete"
                    color = "#38A169"
                else:
                    status = "❌ Failed"
                    color = "#E53E3E"
                
                current_text = item.text()
                # Remove old status if exists
                if "(" in current_text and ")" in current_text:
                    parts = current_text.split(" (")
                    current_text = parts[0]
                
                item.setText(f"{current_text}  {status}")
                item.setForeground(QColor(color))
                break
    # ============================================================
    # VOICE
    # ============================================================
    
    def toggle_voice(self):
        if not self.recognizer or not self.microphone:
            QMessageBox.warning(self, "Voice Not Available", 
                "Voice recognition is not available.\n\nPlease install: pip install SpeechRecognition pyaudio")
            return
        
        self.is_voice_active = not self.is_voice_active
        if self.is_voice_active:
            self._update_status("🎤 Voice active... Speak a command")
            threading.Thread(target=self._listen_voice, daemon=True).start()
        else:
            self._update_status("Voice disabled")
    
    def _listen_voice(self):
        try:
            with self.microphone as source:
                self.recognizer.adjust_for_ambient_noise(source, duration=0.5)
                while self.is_voice_active:
                    try:
                        audio = self.recognizer.listen(source, timeout=5, phrase_time_limit=5)
                        text = self.recognizer.recognize_google(audio)
                        self.voice_command_signal.emit(text.lower())
                    except:
                        continue
        except:
            pass
    
    @Slot(str)
    def _process_voice_command(self, command):
        self._update_status(f"🎤 Voice: {command}")
        response = self.process_voice_command(command)
        self._update_status(f"🤖 {response[:100]}...")
    
    def speak(self, text):
        if self.tts_engine:
            try:
                self.tts_engine.say(text)
                self.tts_engine.runAndWait()
            except:
                pass
    
    def process_voice_command(self, command: str):
        try:
            ai_engine.set_converter(self.converter)
            response = ai_engine.process_request(command, self)
            self.speak(response)
            return response
        except Exception as e:
            logger.error(f"Voice command error: {e}")
            return f"Error processing command: {str(e)}"
    
    # ============================================================
    # CAMERA
    # ============================================================
    
    def toggle_camera(self):
        try:
            import cv2
        except ImportError:
            QMessageBox.warning(self, "Camera Not Available", 
                "OpenCV is not installed.\n\nPlease install: pip install opencv-python")
            return
        
        self.is_camera_active = not self.is_camera_active
        if self.is_camera_active:
            self._update_status("📷 Starting camera...")
            QTimer.singleShot(100, self._scan_with_camera)
        else:
            self._update_status("Camera stopped")
    
    def _scan_with_camera(self):
        """Scan document with camera - Fixed to prevent double saving"""
        try:
            import cv2
            from datetime import datetime
            
            # Create folder for scans
            scan_folder = Path("scanned_documents")
            scan_folder.mkdir(exist_ok=True)
            
            # Open camera
            cap = cv2.VideoCapture(0)
            if not cap.isOpened():
                self._update_status("❌ Camera not accessible")
                self.is_camera_active = False
                if self.camera_btn:
                    self.camera_btn.setText("📷")
                return
            
            # Set camera properties
            cap.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
            cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
            cap.set(cv2.CAP_PROP_FPS, 30)
            
            self._update_status("📷 Press SPACE to capture, ESC to cancel")
            
            window_name = "Document Scanner - Press SPACE to Capture, ESC to Cancel"
            cv2.namedWindow(window_name, cv2.WINDOW_NORMAL)
            cv2.resizeWindow(window_name, 640, 480)
            cv2.moveWindow(window_name, 100, 100)
            
            captured = False
            file_path = None
            scanned_image = None
            
            while self.is_camera_active and cap.isOpened():
                ret, frame = cap.read()
                if not ret:
                    break
                
                overlay = frame.copy()
                h, w = frame.shape[:2]
                cv2.rectangle(overlay, (30, 30), (w-30, h-30), (0, 255, 0), 2)
                cv2.putText(overlay, "SPACE = Capture", (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 255, 0), 2)
                cv2.putText(overlay, "ESC = Cancel", (10, 60), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 0, 255), 2)
                cv2.putText(overlay, "Place document clearly in frame", (10, 90), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 255, 0), 2)
                
                cv2.imshow(window_name, overlay)
                
                key = cv2.waitKey(30) & 0xFF
                
                if key == 27:  # ESC
                    self._update_status("📷 Scan cancelled")
                    captured = True
                    break
                elif key == 32:  # SPACE
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"scanned_document_{timestamp}.jpg"
                    file_path = scan_folder / filename
                    
                    # Save image
                    cv2.imwrite(str(file_path), frame, [cv2.IMWRITE_JPEG_QUALITY, 95])
                    scanned_image = frame
                    
                    self._update_status(f"📷 Document scanned: {filename}")
                    captured = True
                    break
            
            # Cleanup
            cap.release()
            cv2.destroyWindow(window_name)
            
            self.is_camera_active = False
            if self.camera_btn:
                self.camera_btn.setText("📷")
            
            if captured and file_path:
                # ===== FIX: Do NOT add JPG to file list =====
                # Just show the format selection dialog
                self.speak("Document scanned")
                QTimer.singleShot(100, lambda: self._show_format_selection_dialog(str(file_path)))
            else:
                self._update_status("Camera scanner stopped")
                
        except Exception as e:
            logger.error(f"Camera error: {e}")
            self._update_status(f"❌ Camera error: {str(e)[:50]}")
            self.is_camera_active = False
            try:
                cv2.destroyAllWindows()
            except:
                pass
            if self.camera_btn:
                self.camera_btn.setText("📷")
    
    def _show_format_selection_dialog(self, image_path: str):
        """Show dialog to select output format and actions - IMPROVED"""
        from PySide6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QGroupBox, QCheckBox, QComboBox, QFrame, QFileDialog, QMessageBox
        from PySide6.QtCore import Qt
        from pathlib import Path
        from datetime import datetime
        
        dialog = QDialog(self)
        dialog.setWindowTitle("📷 Scan Options")
        dialog.setFixedSize(500, 460)
        dialog.setModal(True)
        
        # ===== FIX: Better dialog styling =====
        dialog.setStyleSheet("""
            QDialog {
                background: #FFFFFF;
                border-radius: 12px;
            }
            QGroupBox {
                border: 1px solid #E2E8F0;
                border-radius: 8px;
                padding-top: 12px;
                margin-top: 4px;
                font-weight: 600;
                font-size: 11px;
                color: #1A202C;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 8px;
                color: #1A202C;
                background: #FFFFFF;
            }
            QPushButton {
                font-size: 11px;
                font-weight: 500;
                padding: 6px 14px;
                border-radius: 6px;
            }
            QPushButton[class="primary"] {
                background: #2E5A8A;
                color: #FFFFFF;
                border: none;
            }
            QPushButton[class="primary"]:hover {
                background: #1E3A5F;
            }
            QPushButton[class="secondary"] {
                background: #EDF2F7;
                color: #2D3748;
                border: 1px solid #E2E8F0;
            }
            QPushButton[class="secondary"]:hover {
                background: #E2E8F0;
            }
            QCheckBox {
                font-size: 11px;
                color: #2D3748;
                spacing: 6px;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
                border-radius: 4px;
                border: 2px solid #CBD5E0;
                background: #FFFFFF;
            }
            QCheckBox::indicator:checked {
                background: #2E5A8A;
                border-color: #2E5A8A;
            }
            QComboBox {
                padding: 4px 8px;
                border: 1px solid #E2E8F0;
                border-radius: 4px;
                font-size: 11px;
                background: #FFFFFF;
                color: #2D3748;
            }
            QComboBox:focus {
                border-color: #2E5A8A;
            }
            QLabel {
                color: #2D3748;
            }
        """)
        
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(20, 16, 20, 16)
        layout.setSpacing(12)
        
        # ===== Header =====
        header = QLabel("📷 Document Scanned Successfully!")
        header.setStyleSheet("""
            font-size: 18px; 
            font-weight: 700; 
            color: #1A202C;
            font-family: 'Segoe UI', sans-serif;
        """)
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(header)
        
        info = QLabel(f"📄 {Path(image_path).name}")
        info.setStyleSheet("""
            font-size: 12px; 
            color: #718096; 
            font-family: 'Segoe UI', sans-serif;
        """)
        info.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(info)
        
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setStyleSheet("background: #E2E8F0; max-height: 1px;")
        layout.addWidget(separator)
        
        # ===== Format selection =====
        format_group = QGroupBox("📄 Convert to Format")
        format_layout = QVBoxLayout(format_group)
        format_layout.setSpacing(8)
        
        format_row1 = QHBoxLayout()
        format_row1.setSpacing(6)
        
        format_buttons = [
            ("📄 PDF", "pdf"),
            ("📝 DOCX", "docx"),
            ("🖼 PNG", "png"),
            ("🖼 JPG", "jpg"),
            ("📊 TIFF", "tiff"),
            ("🌐 WEBP", "webp"),
        ]
        
        self._format_buttons = []
        for label, fmt in format_buttons:
            btn = QPushButton(label)
            btn.setProperty("class", "secondary")
            btn.setCheckable(True)
            btn.setCursor(Qt.PointingHandCursor)
            btn.setStyleSheet("""
                QPushButton {
                    padding: 6px 12px;
                    font-size: 10px;
                    min-width: 55px;
                    border: 2px solid transparent;
                }
                QPushButton:checked {
                    background: #2E5A8A;
                    color: #FFFFFF;
                    border-color: #2E5A8A;
                }
                QPushButton:checked:hover {
                    background: #1E3A5F;
                }
            """)
            btn.clicked.connect(lambda checked, f=fmt: self._select_format(f))
            format_row1.addWidget(btn)
            self._format_buttons.append((btn, fmt))
        
        if self._format_buttons:
            self._format_buttons[0][0].setChecked(True)
            self.selected_format = "pdf"
        
        format_layout.addLayout(format_row1)
        layout.addWidget(format_group)
        
        # ===== OCR option =====
        ocr_group = QGroupBox("📝 Extract Text (OCR)")
        ocr_layout = QVBoxLayout(ocr_group)
        ocr_layout.setSpacing(6)
        
        self.ocr_checkbox = QCheckBox("Extract text using OCR")
        self.ocr_checkbox.setStyleSheet("""
            font-size: 11px; 
            color: #2D3748;
            font-weight: 500;
        """)
        ocr_layout.addWidget(self.ocr_checkbox)
        
        lang_layout = QHBoxLayout()
        lang_layout.setSpacing(8)
        lang_label = QLabel("Language:")
        lang_label.setStyleSheet("font-size: 11px; color: #4A5568;")
        lang_layout.addWidget(lang_label)
        
        self.ocr_lang_combo = QComboBox()
        self.ocr_lang_combo.addItems(['English (en)', 'Urdu (ur)', 'Arabic (ar)'])
        self.ocr_lang_combo.setEnabled(False)
        self.ocr_lang_combo.setStyleSheet("""
            padding: 3px 8px;
            min-width: 100px;
            font-size: 11px;
        """)
        self.ocr_checkbox.toggled.connect(self.ocr_lang_combo.setEnabled)
        lang_layout.addWidget(self.ocr_lang_combo)
        lang_layout.addStretch()
        ocr_layout.addLayout(lang_layout)
        
        layout.addWidget(ocr_group)
        
        # ===== Buttons =====
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)
        
        convert_btn = QPushButton("🔄 Convert & Save")
        convert_btn.setProperty("class", "primary")
        convert_btn.setCursor(Qt.PointingHandCursor)
        convert_btn.setMinimumHeight(36)
        convert_btn.setStyleSheet("""
            QPushButton {
                font-size: 12px;
                font-weight: 600;
                padding: 8px 24px;
            }
        """)
        convert_btn.clicked.connect(lambda: self._process_scanned_document(image_path, dialog))
        btn_layout.addWidget(convert_btn)
        
        cancel_btn = QPushButton("✖ Cancel")
        cancel_btn.setProperty("class", "secondary")
        cancel_btn.setCursor(Qt.PointingHandCursor)
        cancel_btn.setMinimumHeight(36)
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(cancel_btn)
        
        layout.addLayout(btn_layout)
        
        # Execute dialog
        dialog.exec() 

    def _select_format(self, format_name: str):
        """Select a format and unselect others"""
        for btn, fmt in self._format_buttons:
            if fmt == format_name:
                btn.setChecked(True)
            else:
                btn.setChecked(False)
        self.selected_format = format_name

    def _process_scanned_document(self, image_path: str, dialog: QDialog):
        """Process scanned document with selected options"""
        selected_format = None
        for btn, fmt in self._format_buttons:
            if btn.isChecked():
                selected_format = fmt
                break
        
        if not selected_format:
            QMessageBox.warning(self, "No Format", "Please select an output format!")
            return
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"scanned_{timestamp}.{selected_format}"
        
        # Create output directory
        output_dir = Path("scanned_documents")
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / output_filename
        
        try:
            self._update_status(f"🔄 Converting to {selected_format.upper()}...")
            
            from PIL import Image
            import os
            
            # Open the image
            img = Image.open(image_path)
            
            # ===== PDF SAVING =====
            if selected_format == 'pdf':
                # Convert to RGB if needed
                if img.mode == 'RGBA':
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[3])
                    img = background
                elif img.mode == 'P':
                    img = img.convert('RGB')
                elif img.mode == 'LA':
                    img = img.convert('RGB')
                elif img.mode == 'L':
                    img = img.convert('RGB')
                
                # Use reportlab for valid PDF
                try:
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import letter
                    import tempfile
                    
                    c = canvas.Canvas(str(output_path), pagesize=letter)
                    width, height = letter
                    
                    img_width, img_height = img.size
                    aspect = img_width / img_height
                    
                    if aspect > 1:
                        display_width = width - 40
                        display_height = display_width / aspect
                    else:
                        display_height = height - 40
                        display_width = display_height * aspect
                    
                    x = (width - display_width) / 2
                    y = (height - display_height) / 2
                    
                    temp_img = Path(tempfile.gettempdir()) / f"temp_scan_{timestamp}.jpg"
                    img.save(str(temp_img), "JPEG", quality=95)
                    c.drawImage(str(temp_img), x, y, display_width, display_height)
                    c.save()
                    
                    try:
                        temp_img.unlink()
                    except:
                        pass
                        
                except ImportError:
                    img.save(str(output_path), "PDF", resolution=100.0)
            
            elif selected_format in ['jpg', 'jpeg']:
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                img.save(str(output_path), "JPEG", quality=95)
            
            elif selected_format == 'png':
                if img.mode == 'P':
                    img = img.convert('RGBA')
                img.save(str(output_path), "PNG", optimize=True)
            
            elif selected_format == 'webp':
                img.save(str(output_path), "WEBP", quality=90)
            
            elif selected_format == 'tiff':
                img.save(str(output_path), "TIFF", compression='tiff_deflate')
            
            elif selected_format == 'docx':
                from docx import Document
                from docx.shared import Inches
                doc = Document()
                doc.add_heading("Scanned Document", 0)
                doc.add_paragraph(f"File: {Path(image_path).name}")
                doc.add_paragraph(f"Scanned on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.save(str(output_path))
            
            else:
                img.save(str(output_path), "PNG")
            
            # ===== FIX: Only add the converted file ONCE =====
            self._add_file_to_list(str(output_path))
            self.update_file_count()
            
            # ===== FIX: Delete the original JPG after conversion =====
            try:
                original_path = Path(image_path)
                if original_path.exists() and original_path.suffix.lower() == '.jpg':
                    original_path.unlink()
                    logger.info(f"Deleted original JPG: {original_path}")
            except Exception as e:
                logger.warning(f"Could not delete original JPG: {e}")
            
            # OCR if checked
            if hasattr(self, 'ocr_checkbox') and self.ocr_checkbox and self.ocr_checkbox.isChecked():
                threading.Thread(target=self._extract_text_from_image_safe, 
                            args=(image_path, dialog, output_path), daemon=True).start()
            else:
                dialog.accept()
                self._update_status(f"✅ Document saved as: {output_filename}")
                
                # Show success message
                msg_box = QMessageBox(self)
                msg_box.setIcon(QMessageBox.Information)
                msg_box.setWindowTitle("Scan Complete")
                msg_box.setText(f"✅ Document saved as: {output_filename}")
                msg_box.setInformativeText(f"Location: {output_path}")
                
                open_btn = msg_box.addButton("📂 Open File", QMessageBox.AcceptRole)
                close_btn = msg_box.addButton("Close", QMessageBox.RejectRole)
                msg_box.setDefaultButton(close_btn)
                
                msg_box.exec()
                
                if msg_box.clickedButton() == open_btn:
                    self._open_file(str(output_path))
            
        except Exception as e:
            logger.error(f"Processing error: {e}")
            dialog.accept()
            QMessageBox.warning(self, "Error", f"Failed to process document: {str(e)}")
    
    def _extract_text_from_image_safe(self, image_path: str, dialog: QDialog, output_path: Path):
        """Safe OCR extraction in background thread - FIXED for PDF"""
        try:
            text = ""
            
            # For PDF files, extract text directly
            if output_path.suffix.lower() == '.pdf':
                try:
                    import PyPDF2
                    with open(output_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"PDF text extraction failed: {e}")
            
            # For images, use OCR
            if not text or not text.strip():
                try:
                    import pytesseract
                    from PIL import Image
                    
                    # Use the original image for OCR
                    img = Image.open(image_path)
                    text = pytesseract.image_to_string(img)
                    
                except ImportError:
                    logger.warning("Tesseract not available")
                except Exception as e:
                    logger.warning(f"Tesseract OCR failed: {e}")
            
            # If still no text, try EasyOCR
            if not text or not text.strip():
                try:
                    from app.ocr.ocr_engine import ocr_engine
                    success, ocr_text, metadata = ocr_engine.extract_text(str(image_path), 'en')
                    if success and ocr_text:
                        text = ocr_text
                except Exception as e:
                    logger.warning(f"EasyOCR failed: {e}")
            
            # Save extracted text if found
            if text and text.strip():
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                text_path = Path("scanned_documents") / f"extracted_text_{timestamp}.txt"
                text_path.parent.mkdir(exist_ok=True)
                
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                
                self._add_file_to_list(str(text_path))
                
                dialog.accept()
                
                QMessageBox.information(
                    self,
                    "Scan & OCR Complete",
                    f"✅ Document and text extracted!\n\n"
                    f"📄 Document: {output_path.name}\n"
                    f"📝 Text: {text_path.name}\n"
                    f"📊 Words: {len(text.split())}"
                )
                return
            else:
                dialog.accept()
                QMessageBox.information(
                    self,
                    "Scan Complete",
                    f"✅ Document saved successfully!\n\n"
                    f"📄 File: {output_path.name}\n"
                    f"ℹ️ No text detected in the image."
                )
                
        except Exception as e:
            logger.error(f"OCR error: {e}")
            dialog.accept()

    def _open_file(self, file_path: str):
        """Open a file with the default application"""
        try:
            import os
            import subprocess
            
            if not Path(file_path).exists():
                QMessageBox.warning(self, "File Not Found", f"File not found:\n{file_path}")
                return
            
            # Open with default application
            if os.name == 'nt':  # Windows
                os.startfile(file_path)
            else:  # Mac/Linux
                subprocess.Popen(['xdg-open', file_path])
                
            self._update_status(f"📂 Opened: {Path(file_path).name}")
            
        except Exception as e:
            logger.error(f"Open file error: {e}")
            QMessageBox.warning(self, "Open Error", f"Could not open file:\n{str(e)}")
    # ============================================================
    # OCR
    # ============================================================
    
    def open_ocr(self):
        try:
            from app.gui.ocr_dialog import OCRDialog
            dialog = OCRDialog(self)
            dialog.exec()
        except Exception as e:
            QMessageBox.warning(self, "OCR Error", f"Error opening OCR: {str(e)}")
    
    # ============================================================
    # AI POPUP
    # ============================================================
    
    def open_ai_popup(self):
        """Open AI Assistant as a popup window"""
        try:
            from app.gui.ai_chat import AIChatWidget
            
            popup = QDialog(self)
            popup.setWindowTitle("🤖 AI Assistant")
            popup.setMinimumSize(600, 500)
            popup.resize(650, 550)
            popup.setModal(False)
            popup.setStyleSheet("""
                QDialog {
                    background: #FFFFFF;
                    border-radius: 10px;
                }
            """)
            
            layout = QVBoxLayout(popup)
            layout.setContentsMargins(0, 0, 0, 0)
            layout.setSpacing(0)
            
            # Header
            header = QFrame()
            header.setFixedHeight(50)
            header.setStyleSheet("""
                QFrame {
                    background: #2E5A8A;
                    border-radius: 10px 10px 0 0;
                }
            """)
            header_layout = QHBoxLayout(header)
            header_layout.setContentsMargins(20, 0, 20, 0)
            
            title = QLabel("🤖 AI Assistant")
            title.setStyleSheet("font-size: 17px; font-weight: 700; color: #FFFFFF;")
            header_layout.addWidget(title)
            header_layout.addStretch()
            
            close_btn = QPushButton("✕")
            close_btn.setFixedSize(30, 30)
            close_btn.setStyleSheet("""
                QPushButton {
                    border: none;
                    border-radius: 15px;
                    background: rgba(255,255,255,0.15);
                    color: #FFFFFF;
                    font-size: 14px;
                    font-weight: 600;
                }
                QPushButton:hover {
                    background: rgba(255,255,255,0.25);
                }
            """)
            close_btn.clicked.connect(popup.close)
            header_layout.addWidget(close_btn)
            
            layout.addWidget(header)
            
            # ===== AI CHAT WIDGET - With logging =====
            logger.info("Creating AI Chat Widget...")
            ai_chat = AIChatWidget(popup)
            layout.addWidget(ai_chat, 1)
            logger.info("AI Chat Widget created successfully")
            
            # Show popup
            logger.info("Showing AI popup...")
            popup.exec_()
            logger.info("AI popup closed")
            
        except Exception as e:
            logger.error(f"Error opening AI popup: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.warning(self, "Error", f"Could not open AI Assistant:\n{str(e)}")
            
    # ============================================================
    # SEARCH AND ADD FILE
    # ============================================================
    
    def search_and_add_file(self, file_pattern: str) -> str:
        try:
            search_locations = [
                Path.cwd(),
                Path.home() / "Documents",
                Path.home() / "Downloads",
                Path.home() / "Desktop",
            ]
            
            for location in search_locations:
                if not location.exists():
                    continue
                for file_path in location.glob(f"*{file_pattern}*"):
                    if file_path.is_file():
                        self._add_file_to_list(str(file_path))
                        return f"✅ Found and added: {file_path.name}"
            
            for location in search_locations:
                if not location.exists():
                    continue
                for file_path in location.iterdir():
                    if file_path.is_file() and file_pattern.lower() in file_path.name.lower():
                        self._add_file_to_list(str(file_path))
                        return f"✅ Found and added: {file_path.name}"
            
            return None
        except Exception as e:
            logger.error(f"Error searching file: {e}")
            return None