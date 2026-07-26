"""
E-Book Converter - Complete with Calibre & Pandoc support
Supports: EPUB, MOBI, AZW3, FB2, PDF, DOCX, TXT, HTML, LIT, LRF
"""

from pathlib import Path
import os
import subprocess
import shutil
import tempfile
import re

from app.utils.logger import get_logger

logger = get_logger(__name__)


class EBookConverter:
    """Complete e-book converter with multiple backends"""
    
    def __init__(self):
        # All supported formats
        self.supported_formats = [
            'epub', 'mobi', 'azw3', 'fb2', 'lit', 'lrf',
            'pdf', 'docx', 'txt', 'html', 'rtf', 'odt',
            'azw', 'kfx', 'pdb', 'prc', 'pml', 'rb'
        ]
        
        # Output formats
        self.output_formats = [
            'epub', 'mobi', 'azw3', 'fb2', 'pdf', 'docx',
            'txt', 'html', 'rtf', 'odt'
        ]
        
        # Cache Calibre path
        self._calibre_path = None
        self._pandoc_path = None
        
    def convert(self, input_path: str, output_path: str, options: dict = None):
        """Convert e-book to any supported format"""
        try:
            # Check if input file exists
            if not os.path.exists(input_path):
                return False, f"Input file not found: {input_path}"
            
            input_ext = Path(input_path).suffix[1:].lower()
            output_ext = Path(output_path).suffix[1:].lower()
            
            logger.info(f"Converting e-book: {Path(input_path).name} ({input_ext} -> {output_ext})")
            
            # ===== TRY CALIBRE FIRST (best quality) =====
            if self._is_calibre_available():
                success, message = self._convert_with_calibre(input_path, output_path, options)
                if success:
                    return True, message
            
            # ===== TRY PANDOC (good for markdown/text formats) =====
            if self._is_pandoc_available():
                success, message = self._convert_with_pandoc(input_path, output_path, options)
                if success:
                    return True, message
            
            # ===== FALLBACK: Use Python libraries =====
            # EPUB → other formats
            if input_ext == 'epub':
                if output_ext in ['txt', 'html']:
                    return self._epub_to_text(input_path, output_path, output_ext)
                elif output_ext == 'docx':
                    return self._epub_to_docx(input_path, output_path)
                elif output_ext == 'pdf':
                    return self._epub_to_pdf(input_path, output_path)
            
            # MOBI → other formats
            if input_ext in ['mobi', 'azw3']:
                if output_ext == 'txt':
                    return self._mobi_to_text(input_path, output_path)
                elif output_ext == 'epub':
                    return self._mobi_to_epub(input_path, output_path)
            
            # FB2 → other formats
            if input_ext == 'fb2':
                if output_ext in ['txt', 'html']:
                    return self._fb2_to_text(input_path, output_path, output_ext)
            
            # PDF → EPUB
            if input_ext == 'pdf' and output_ext == 'epub':
                return self._pdf_to_epub(input_path, output_path)
            
            # TXT → EPUB
            if input_ext == 'txt' and output_ext == 'epub':
                return self._txt_to_epub(input_path, output_path, options)
            
            # HTML → EPUB
            if input_ext == 'html' and output_ext == 'epub':
                return self._html_to_epub(input_path, output_path)
            
            # ===== LAST RESORT: Copy =====
            shutil.copy2(input_path, output_path)
            return True, f"E-book copied (no conversion available for {input_ext} → {output_ext})"
            
        except Exception as e:
            logger.error(f"E-book conversion error: {e}")
            return False, f"E-book conversion error: {str(e)}"
    
    # ============================================================
    # CALIBRE CONVERSION (Best Quality)
    # ============================================================
    
    def _is_calibre_available(self) -> bool:
        """Check if Calibre's ebook-convert is available"""
        if self._calibre_path is not None:
            return self._calibre_path is not False
        
        # Check common paths
        common_paths = [
            "C:\\Program Files\\Calibre2\\ebook-convert.exe",
            "C:\\Program Files (x86)\\Calibre2\\ebook-convert.exe",
            os.path.expanduser("~\\AppData\\Local\\calibre\\ebook-convert.exe"),
            "ebook-convert"  # Check PATH
        ]
        
        for path in common_paths:
            try:
                if path == "ebook-convert":
                    result = subprocess.run(['ebook-convert', '--version'],
                                          capture_output=True, timeout=2)
                    if result.returncode == 0:
                        self._calibre_path = "ebook-convert"
                        logger.info("✅ Calibre found in PATH")
                        return True
                elif os.path.exists(path):
                    self._calibre_path = path
                    logger.info(f"✅ Calibre found at: {path}")
                    return True
            except:
                pass
        
        self._calibre_path = False
        logger.warning("⚠️ Calibre not found. Install Calibre for better e-book conversion.")
        return False
    
    def _convert_with_calibre(self, input_path: str, output_path: str, options: dict = None):
        """Convert using Calibre's ebook-convert"""
        try:
            if not self._calibre_path:
                return False, "Calibre not available"
            
            logger.info("Using Calibre for e-book conversion")
            
            # Build command
            cmd = [self._calibre_path, input_path, output_path]
            
            # Add options
            if options:
                if options.get('title'):
                    cmd.extend(['--title', options['title']])
                if options.get('author'):
                    cmd.extend(['--authors', options['author']])
                if options.get('language'):
                    cmd.extend(['--language', options['language']])
                if 'cover' in options and options['cover']:
                    cmd.extend(['--cover', options['cover']])
                if options.get('output_profile'):
                    cmd.extend(['--output-profile', options['output_profile']])
            
            # Run conversion
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if result.returncode == 0 and Path(output_path).exists():
                return True, f"E-book converted with Calibre: {Path(output_path).name}"
            else:
                error = result.stderr or "Unknown error"
                return False, f"Calibre error: {error[:200]}"
            
        except subprocess.TimeoutExpired:
            return False, "Calibre conversion timed out"
        except Exception as e:
            return False, f"Calibre error: {str(e)}"
    
    # ============================================================
    # PANDOC CONVERSION (Good for Text/Markdown)
    # ============================================================
    
    def _is_pandoc_available(self) -> bool:
        """Check if Pandoc is available"""
        if self._pandoc_path is not None:
            return self._pandoc_path is not False
        
        try:
            result = subprocess.run(['pandoc', '--version'],
                                  capture_output=True, timeout=2)
            if result.returncode == 0:
                self._pandoc_path = "pandoc"
                logger.info("✅ Pandoc found in PATH")
                return True
        except:
            pass
        
        self._pandoc_path = False
        return False
    
    def _convert_with_pandoc(self, input_path: str, output_path: str, options: dict = None):
        """Convert using Pandoc"""
        try:
            if not self._pandoc_path:
                return False, "Pandoc not available"
            
            logger.info("Using Pandoc for e-book conversion")
            
            cmd = [
                self._pandoc_path,
                input_path,
                '-o', output_path
            ]
            
            # Add options
            if options:
                if options.get('title'):
                    cmd.extend(['--metadata', f'title={options["title"]}'])
                if options.get('author'):
                    cmd.extend(['--metadata', f'author={options["author"]}'])
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and Path(output_path).exists():
                return True, f"E-book converted with Pandoc: {Path(output_path).name}"
            else:
                return False, f"Pandoc error: {result.stderr[:200]}"
            
        except Exception as e:
            return False, f"Pandoc error: {str(e)}"
    
    # ============================================================
    # EPUB CONVERSIONS (Python-based)
    # ============================================================
    
    def _epub_to_text(self, input_path: str, output_path: str, output_ext: str):
        """Convert EPUB to TXT or HTML using zipfile"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            content = []
            
            with zipfile.ZipFile(input_path, 'r') as epub:
                # Find content files
                content_files = []
                for name in epub.namelist():
                    if name.endswith(('.xhtml', '.html', '.xml')):
                        content_files.append(name)
                    elif 'content.opf' in name:
                        # Parse OPF to get spine order
                        try:
                            with epub.open(name) as f:
                                tree = ET.parse(f)
                                root = tree.getroot()
                                ns = {'opf': 'http://www.idpf.org/2007/opf'}
                                for item in root.findall('.//opf:item', ns):
                                    if item.get('media-type', '').startswith('application/xhtml'):
                                        href = item.get('href')
                                        if href:
                                            # Get the file path
                                            base_dir = Path(name).parent
                                            content_files.append(str(base_dir / href))
                        except:
                            pass
            
            # Extract text from content files
            for file_name in content_files:
                try:
                    with epub.open(file_name) as f:
                        data = f.read().decode('utf-8', errors='ignore')
                        # Remove HTML tags
                        text = re.sub(r'<[^>]+>', ' ', data)
                        text = re.sub(r'\s+', ' ', text).strip()
                        if text:
                            content.append(text)
                except:
                    pass
            
            if not content:
                return False, "No text content found in EPUB"
            
            full_text = '\n\n'.join(content)
            
            if output_ext == 'html':
                html = ['<!DOCTYPE html>', '<html><head><meta charset="utf-8">',
                        '<title>EPUB Content</title></head><body>']
                for para in content:
                    html.append(f'<p>{para}</p>')
                html.append('</body></html>')
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(html))
            else:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(full_text)
            
            return True, f"EPUB to {output_ext.upper()} conversion successful"
            
        except Exception as e:
            return False, f"EPUB to text error: {str(e)}"
    
    def _epub_to_docx(self, input_path: str, output_path: str):
        """Convert EPUB to DOCX using docx library"""
        try:
            from docx import Document
            from docx.shared import Pt
            
            import zipfile
            import re
            
            doc = Document()
            
            with zipfile.ZipFile(input_path, 'r') as epub:
                for name in epub.namelist():
                    if name.endswith(('.xhtml', '.html', '.xml')):
                        try:
                            with epub.open(name) as f:
                                data = f.read().decode('utf-8', errors='ignore')
                                text = re.sub(r'<[^>]+>', ' ', data)
                                text = re.sub(r'\s+', ' ', text).strip()
                                if text:
                                    doc.add_paragraph(text)
                        except:
                            pass
            
            doc.save(output_path)
            return True, "EPUB to DOCX conversion successful"
            
        except Exception as e:
            return False, f"EPUB to DOCX error: {str(e)}"
    
    def _epub_to_pdf(self, input_path: str, output_path: str):
        """Convert EPUB to PDF using reportlab"""
        try:
            import zipfile
            import re
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.utils import simpleSplit
            
            # Extract text first
            content = []
            with zipfile.ZipFile(input_path, 'r') as epub:
                for name in epub.namelist():
                    if name.endswith(('.xhtml', '.html', '.xml')):
                        try:
                            with epub.open(name) as f:
                                data = f.read().decode('utf-8', errors='ignore')
                                text = re.sub(r'<[^>]+>', ' ', data)
                                text = re.sub(r'\s+', ' ', text).strip()
                                if text:
                                    content.append(text)
                        except:
                            pass
            
            if not content:
                return False, "No text content found in EPUB"
            
            # Create PDF
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            y = height - 50
            line_height = 14
            font_size = 11
            
            c.setFont("Helvetica", font_size)
            
            for text in content:
                words = text.split()
                line = []
                for word in words:
                    test_line = ' '.join(line + [word])
                    if c.stringWidth(test_line, "Helvetica", font_size) < width - 100:
                        line.append(word)
                    else:
                        if y < 50:
                            c.showPage()
                            c.setFont("Helvetica", font_size)
                            y = height - 50
                        c.drawString(50, y, ' '.join(line))
                        y -= line_height
                        line = [word]
                
                if line:
                    if y < 50:
                        c.showPage()
                        c.setFont("Helvetica", font_size)
                        y = height - 50
                    c.drawString(50, y, ' '.join(line))
                    y -= line_height
                
                # Add spacing between paragraphs
                y -= line_height
            
            c.save()
            return True, "EPUB to PDF conversion successful"
            
        except Exception as e:
            return False, f"EPUB to PDF error: {str(e)}"
    
    # ============================================================
    # MOBI/AZW3 CONVERSIONS
    # ============================================================
    
    def _mobi_to_text(self, input_path: str, output_path: str):
        """Extract text from MOBI/AZW3 (simplified)"""
        try:
            # Use Calibre if available
            if self._is_calibre_available():
                temp_file = tempfile.NamedTemporaryFile(suffix='.txt', delete=False)
                temp_path = temp_file.name
                temp_file.close()
                
                success, _ = self._convert_with_calibre(input_path, temp_path, None)
                if success:
                    shutil.move(temp_path, output_path)
                    return True, "MOBI to TXT conversion successful"
            
            # Fallback: Use mobi library if available
            try:
                import mobi
                temp_dir = tempfile.mkdtemp()
                mobi.extract(input_path, temp_dir)
                
                # Find text files
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        if file.endswith(('.html', '.htm', '.txt')):
                            with open(os.path.join(root, file), 'r', encoding='utf-8', errors='ignore') as f:
                                content = f.read()
                            with open(output_path, 'w', encoding='utf-8') as f:
                                f.write(content)
                            shutil.rmtree(temp_dir)
                            return True, "MOBI to TXT conversion successful"
                
                shutil.rmtree(temp_dir)
            except ImportError:
                pass
            except Exception as e:
                logger.warning(f"MOBI extraction failed: {e}")
            
            return False, "Could not extract text from MOBI. Install Calibre for better support."
            
        except Exception as e:
            return False, f"MOBI to text error: {str(e)}"
    
    def _mobi_to_epub(self, input_path: str, output_path: str):
        """Convert MOBI to EPUB using Calibre"""
        try:
            if self._is_calibre_available():
                return self._convert_with_calibre(input_path, output_path, None)
            
            # Try using kindleunpack if available
            try:
                import kindleunpack
                # kindleunpack extraction
                temp_dir = tempfile.mkdtemp()
                kindleunpack.unpack_book(input_path, temp_dir)
                
                # Re-pack as EPUB (simplified)
                # This is complex; recommend Calibre
                shutil.rmtree(temp_dir)
            except ImportError:
                pass
            
            return False, "MOBI to EPUB requires Calibre. Please install Calibre."
            
        except Exception as e:
            return False, f"MOBI to EPUB error: {str(e)}"
    
    # ============================================================
    # FB2 CONVERSIONS
    # ============================================================
    
    def _fb2_to_text(self, input_path: str, output_path: str, output_ext: str):
        """Extract text from FB2"""
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(input_path)
            root = tree.getroot()
            
            # Find all paragraphs
            ns = {'fb': 'http://www.gribuser.ru/xml/fictionbook/2.0'}
            paragraphs = root.findall('.//fb:p', ns)
            
            content = []
            for p in paragraphs:
                if p.text:
                    content.append(p.text.strip())
            
            if not content:
                return False, "No text found in FB2"
            
            full_text = '\n\n'.join(content)
            
            if output_ext == 'html':
                html = ['<!DOCTYPE html>', '<html><head><meta charset="utf-8">',
                        '<title>FB2 Content</title></head><body>']
                for para in content:
                    html.append(f'<p>{para}</p>')
                html.append('</body></html>')
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(html))
            else:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(full_text)
            
            return True, f"FB2 to {output_ext.upper()} conversion successful"
            
        except Exception as e:
            return False, f"FB2 to text error: {str(e)}"
    
    # ============================================================
    # OTHER CONVERSIONS
    # ============================================================
    
    def _pdf_to_epub(self, input_path: str, output_path: str):
        """Convert PDF to EPUB"""
        try:
            if self._is_calibre_available():
                return self._convert_with_calibre(input_path, output_path, None)
            
            # Use PyPDF2 to extract text
            import PyPDF2
            from docx import Document
            
            doc = Document()
            with open(input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
            
            # Save as temp docx then convert
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_path = temp_docx.name
            temp_docx.close()
            doc.save(temp_path)
            
            # Convert DOCX to EPUB using Calibre or Pandoc
            if self._is_calibre_available():
                success, msg = self._convert_with_calibre(temp_path, output_path, None)
                os.unlink(temp_path)
                return success, msg
            
            os.unlink(temp_path)
            return False, "PDF to EPUB requires Calibre. Please install Calibre."
            
        except Exception as e:
            return False, f"PDF to EPUB error: {str(e)}"
    
    def _txt_to_epub(self, input_path: str, output_path: str, options: dict = None):
        """Convert TXT to EPUB"""
        try:
            if self._is_pandoc_available():
                return self._convert_with_pandoc(input_path, output_path, options)
            
            if self._is_calibre_available():
                return self._convert_with_calibre(input_path, output_path, options)
            
            # Fallback: Create a simple EPUB
            import zipfile
            from datetime import datetime
            
            title = options.get('title', Path(input_path).stem) if options else Path(input_path).stem
            
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Create EPUB structure
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_dir = Path(temp_dir)
                (temp_dir / 'META-INF').mkdir(parents=True)
                (temp_dir / 'OEBPS').mkdir(parents=True)
                
                # META-INF/container.xml
                container = '''<?xml version="1.0" encoding="UTF-8"?>
                <container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
                  <rootfiles>
                    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
                  </rootfiles>
                </container>'''
                with open(temp_dir / 'META-INF' / 'container.xml', 'w') as f:
                    f.write(container)
                
                # OEBPS/content.opf
                opf = f'''<?xml version="1.0" encoding="UTF-8"?>
                <package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="BookId">
                  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
                    <dc:title>{title}</dc:title>
                    <dc:language>en</dc:language>
                    <dc:date>{datetime.now().isoformat()}</dc:date>
                  </metadata>
                  <manifest>
                    <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
                    <item id="text" href="text.xhtml" media-type="application/xhtml+xml"/>
                  </manifest>
                  <spine>
                    <itemref idref="text"/>
                  </spine>
                </package>'''
                with open(temp_dir / 'OEBPS' / 'content.opf', 'w') as f:
                    f.write(opf)
                
                # OEBPS/toc.ncx
                ncx = f'''<?xml version="1.0" encoding="UTF-8"?>
                <ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1" xml:lang="en">
                  <head>
                    <meta name="dtb:uid" content="BookId"/>
                    <meta name="dtb:depth" content="1"/>
                    <meta name="dtb:totalPageCount" content="0"/>
                    <meta name="dtb:maxPageNumber" content="0"/>
                  </head>
                  <docTitle><text>{title}</text></docTitle>
                  <navMap><navPoint id="navpoint-1" playOrder="1">
                    <navLabel><text>Start</text></navLabel>
                    <content src="text.xhtml"/>
                  </navPoint></navMap>
                </ncx>'''
                with open(temp_dir / 'OEBPS' / 'toc.ncx', 'w') as f:
                    f.write(ncx)
                
                # OEBPS/text.xhtml
                xhtml = f'''<?xml version="1.0" encoding="UTF-8"?>
                <!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
                <html xmlns="http://www.w3.org/1999/xhtml">
                <head><title>{title}</title></head>
                <body>'''
                
                for line in content.split('\n'):
                    if line.strip():
                        xhtml += f'<p>{line}</p>'
                
                xhtml += '</body></html>'
                with open(temp_dir / 'OEBPS' / 'text.xhtml', 'w', encoding='utf-8') as f:
                    f.write(xhtml)
                
                # Create EPUB
                with zipfile.ZipFile(output_path, 'w') as epub:
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = Path(root) / file
                            arc_name = str(file_path.relative_to(temp_dir))
                            epub.write(file_path, arc_name)
            
            return True, "TXT to EPUB conversion successful"
            
        except Exception as e:
            return False, f"TXT to EPUB error: {str(e)}"
    
    def _html_to_epub(self, input_path: str, output_path: str):
        """Convert HTML to EPUB"""
        try:
            if self._is_pandoc_available():
                return self._convert_with_pandoc(input_path, output_path, None)
            
            if self._is_calibre_available():
                return self._convert_with_calibre(input_path, output_path, None)
            
            return False, "HTML to EPUB requires Pandoc or Calibre"
            
        except Exception as e:
            return False, f"HTML to EPUB error: {str(e)}"
    
    # ============================================================
    # HELPER METHODS
    # ============================================================
    
    def get_supported_formats(self) -> list:
        """Get all supported input formats"""
        return self.supported_formats
    
    def get_output_formats(self) -> list:
        """Get all supported output formats"""
        return self.output_formats
    
    def get_format_info(self, format_name: str) -> dict:
        """Get information about a format"""
        info = {
            'epub': {'name': 'EPUB', 'extension': 'epub', 'mime': 'application/epub+zip'},
            'mobi': {'name': 'MOBI', 'extension': 'mobi', 'mime': 'application/x-mobipocket-ebook'},
            'azw3': {'name': 'AZW3', 'extension': 'azw3', 'mime': 'application/vnd.amazon.ebook'},
            'fb2': {'name': 'FB2', 'extension': 'fb2', 'mime': 'text/fb2+xml'},
            'pdf': {'name': 'PDF', 'extension': 'pdf', 'mime': 'application/pdf'},
            'docx': {'name': 'DOCX', 'extension': 'docx', 'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
            'txt': {'name': 'TXT', 'extension': 'txt', 'mime': 'text/plain'},
            'html': {'name': 'HTML', 'extension': 'html', 'mime': 'text/html'},
            'rtf': {'name': 'RTF', 'extension': 'rtf', 'mime': 'text/rtf'},
            'odt': {'name': 'ODT', 'extension': 'odt', 'mime': 'application/vnd.oasis.opendocument.text'},
            'lit': {'name': 'LIT', 'extension': 'lit', 'mime': 'application/vnd.ms-ebook'},
            'lrf': {'name': 'LRF', 'extension': 'lrf', 'mime': 'application/x-sony-bbeb'},
            'azw': {'name': 'AZW', 'extension': 'azw', 'mime': 'application/vnd.amazon.ebook'},
            'pdb': {'name': 'PDB', 'extension': 'pdb', 'mime': 'application/vnd.palm'},
            'prc': {'name': 'PRC', 'extension': 'prc', 'mime': 'application/x-mobipocket-ebook'},
        }
        return info.get(format_name.lower(), {})
    
    def get_conversion_requirements(self, input_format: str, output_format: str) -> dict:
        """Get requirements for a specific conversion"""
        input_ext = input_format.lower()
        output_ext = output_format.lower()
        
        requirements = {
            'calibre': ['epub', 'mobi', 'azw3', 'fb2', 'pdf', 'docx', 'txt', 'html', 'rtf', 'odt'],
            'pandoc': ['epub', 'docx', 'txt', 'html', 'rtf', 'odt', 'pdf'],
        }
        
        result = {
            'requires_calibre': False,
            'requires_pandoc': False,
            'requires_native': False
        }
        
        # Check if Calibre is needed
        if input_ext in ['mobi', 'azw3', 'lit', 'lrf'] and output_ext in ['epub', 'pdf']:
            result['requires_calibre'] = True
        
        # Check if Pandoc is needed
        if input_ext in ['md', 'markdown'] and output_ext in ['epub', 'docx']:
            result['requires_pandoc'] = True
        
        return result